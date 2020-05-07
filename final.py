"""
Laura Pankow
This program creates a taco cookbook with 3 randomly generated taco recipes.
"""
import image as image
import requests
import docx
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt

# creating a variable for the url that generates random taco recipes
url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'

# creating three recipe variables, all getting info from the url
recipe1 = requests.get(url).json()
recipe2 = requests.get(url).json()
recipe3 = requests.get(url).json()

# creating a variable for recipe 1, a string that will be the title of the recipe, pulling info from the recipe index
recipe1String = recipe1['base_layer']['name'] + ' with ' + recipe1['mixin']['name'] + ' and ' + recipe1['condiment']['name']
# creating variables that will contain the recipe for each part of the taco, grabbing positions from the json data keys
seasoningRecipe1 = recipe1['seasoning']['recipe']
condimentRecipe1 = recipe1['condiment']['recipe']
mixinRecipe1 = recipe1 ['mixin']['recipe']
base_layerRecipe1 = recipe1['base_layer']['recipe']
shellRecipe1 = recipe1['shell']['recipe']

# creating a variable for recipe 2, a string that will be the title of the recipe, pulling info from the recipe index
recipe2String = recipe2['base_layer']['name'] + ' with ' + recipe2['mixin']['name'] + ' and ' + recipe2['condiment']['name']
# creating variables that will contain the recipe for each part of the taco, grabbing positions from the json data keys
seasoningRecipe2 = recipe2['seasoning']['recipe']
condimentRecipe2 = recipe2['condiment']['recipe']
mixinRecipe2 = recipe2['mixin']['recipe']
base_layerRecipe2 = recipe2['base_layer']['recipe']
shellRecipe2 = recipe2['shell']['recipe']

# creating a variable for recipe 3, a string that will be the title of the recipe, pulling info from the recipe index
recipe3String = recipe3['base_layer']['name'] + ' with ' + recipe3['mixin']['name'] + ' and ' + recipe3['condiment']['name']
# creating variables that will contain the recipe for each part of the taco, grabbing positions from the json data keys
seasoningRecipe3= recipe3['seasoning']['recipe']
condimentRecipe3 = recipe3['condiment']['recipe']
mixinRecipe3 = recipe3['mixin']['recipe']
base_layerRecipe3 = recipe3['base_layer']['recipe']
shellRecipe3 = recipe3['shell']['recipe']

# the following image processes use the pillow library
# creating a variable to open the taco pic, resizing the pic, then saving the resized pic
image = Image.open('tacopic.jpg')
image = image.resize((400,400))
image.save('tacopic.jpg')

# opening the taco pic again, using the ImageDraw to write on the picture, setting the font, font size, color, and location of the text, then saving the edited pic
image = Image.open('tacopic.jpg')
imgDraw = ImageDraw.Draw(image)
font = ImageFont.truetype('DejaVuSans.ttf', 20)
imgDraw.text([0, 350], 'Random Taco Cookbook', fill='black', font=font)
image.save('tacopic2.jpg')

# creating a new word document using the docx library
document = docx.Document()
# adding a heading (title), on the first line, and setting the font size
newHeading = document.add_heading(level=0)
writeHeading = newHeading.add_run('Random Taco Cookbook')
writeHeading.font.size = Pt(24)

# adding the edited taco picture to the title page
document.add_picture('tacopic2.jpg')

# adding a paragraph, with the Heading 1 format to make stand out. Then adding another paragraph with credit info
document.add_paragraph('Credits:\n', 'Heading 1')
document.add_paragraph(f'Taco Image: Photo by Tai\'s Captures on Unsplash\n' 
                       'Taco Recipes From: https://taco-1150.herokuapp.com/random/?full_taco=true\n'
                       'Code By: Laura Pankow\n')
# creating a page break
document.add_page_break()

# adding a new heading (title) for the new page on the first line
newHeading = document.add_heading(level=0)
# new heading using the variable for the recipe 1 title string, setting the font size
writeHeading = newHeading.add_run(f'{recipe1String}')
writeHeading.font.size = Pt(24)
# adding new paragraphs of each recipe name, formatted to Heading 1 so that they stand out, pulling info from index
document.add_paragraph(f'{recipe1["seasoning"]["name"]}\n', 'Heading 1')

# the following .add_paragraphs are adding the recipes below there titles, using the variables created above, for each section
document.add_paragraph(f'{seasoningRecipe1}\n')
document.add_paragraph(f'{recipe1["condiment"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{condimentRecipe1}\n')
document.add_paragraph(f'{recipe1["mixin"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{mixinRecipe1}\n')
document.add_paragraph(f'{recipe1["base_layer"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{base_layerRecipe1}\n')
document.add_paragraph(f'{recipe1["shell"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{shellRecipe1}')
document.add_page_break()

p = document.add_heading(level=0)
wp = p.add_run(f'{recipe2String}')
wp.font.size = Pt(24)
document.add_paragraph(f'{recipe2["seasoning"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{seasoningRecipe2}\n')
document.add_paragraph(f'{recipe2["condiment"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{condimentRecipe2}\n')
document.add_paragraph(f'{recipe2["mixin"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{mixinRecipe2}\n')
document.add_paragraph(f'{recipe2["base_layer"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{base_layerRecipe2}\n')
document.add_paragraph(f'{recipe2["shell"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{shellRecipe2}')
document.add_page_break()

p = document.add_heading(level=0)
wp = p.add_run(f'{recipe3String}')
wp.font.size = Pt(24)
document.add_paragraph(f'{recipe3["seasoning"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{seasoningRecipe3}\n')
document.add_paragraph(f'{recipe3["condiment"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{condimentRecipe3}\n')
document.add_paragraph(f'{recipe3["mixin"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{mixinRecipe3}\n')
document.add_paragraph(f'{recipe3["base_layer"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{base_layerRecipe3}\n')
document.add_paragraph(f'{recipe3["shell"]["name"]}\n', 'Heading 1')
document.add_paragraph(f'{shellRecipe3}')

# saving the word document
document.save('RandomTacoCookbook.docx')